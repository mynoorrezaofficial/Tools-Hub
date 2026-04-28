import os
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib import colors
from docx import Document
from docx.shared import Inches, Pt

def generate_cv_pdf(data, output_path, template='classic'):
    """
    Generates a CV in PDF format based on provided data and template.
    """
    doc = SimpleDocTemplate(output_path, pagesize=letter)
    styles = getSampleStyleSheet()
    story = []

    # Custom Styles
    name_style = ParagraphStyle(
        'NameStyle',
        parent=styles['Heading1'],
        fontSize=24,
        spaceAfter=12,
        textColor=colors.HexColor("#1e293b")
    )
    
    section_title_style = ParagraphStyle(
        'SectionTitle',
        parent=styles['Heading2'],
        fontSize=14,
        spaceBefore=12,
        spaceAfter=6,
        textColor=colors.HexColor("#2563eb"),
        borderPadding=2,
        borderWidth=0,
        borderColor=colors.white,
        backColor=colors.white
    )

    # Personal Info
    personal = data.get('personal', {})
    story.append(Paragraph(personal.get('name', 'Your Name'), name_style))
    contact_info = f"{personal.get('email', '')} | {personal.get('phone', '')} | {personal.get('location', '')}"
    story.append(Paragraph(contact_info, styles['Normal']))
    if personal.get('linkedin'):
        story.append(Paragraph(f"LinkedIn: {personal.get('linkedin')}", styles['Normal']))
    
    story.append(Spacer(1, 12))

    # Summary
    if data.get('summary'):
        story.append(Paragraph("Professional Summary", section_title_style))
        story.append(Paragraph(data['summary'], styles['Normal']))
        story.append(Spacer(1, 12))

    # Work Experience
    if data.get('experience'):
        story.append(Paragraph("Work Experience", section_title_style))
        for exp in data['experience']:
            title = f"<b>{exp.get('role', '')}</b> at {exp.get('company', '')}"
            date = f"{exp.get('startDate', '')} - {exp.get('endDate', 'Present')}"
            story.append(Paragraph(f"{title} ({date})", styles['Normal']))
            story.append(Paragraph(exp.get('description', ''), styles['Normal']))
            story.append(Spacer(1, 8))

    # Education
    if data.get('education'):
        story.append(Paragraph("Education", section_title_style))
        for edu in data['education']:
            degree = f"<b>{edu.get('degree', '')}</b>, {edu.get('school', '')}"
            date = f"{edu.get('startDate', '')} - {edu.get('endDate', '')}"
            story.append(Paragraph(f"{degree} ({date})", styles['Normal']))
            story.append(Spacer(1, 4))

    # Skills
    if data.get('skills'):
        story.append(Paragraph("Skills", section_title_style))
        skills_text = ", ".join(data['skills'])
        story.append(Paragraph(skills_text, styles['Normal']))

    doc.build(story)
    return True, "PDF Generated Successfully"

def generate_cv_docx(data, output_path, template='classic'):
    """
    Generates a CV in DOCX format based on provided data and template.
    """
    doc = Document()
    
    # Personal Info
    personal = data.get('personal', {})
    name = doc.add_heading(personal.get('name', 'Your Name'), 0)
    
    contact_info = f"{personal.get('email', '')} | {personal.get('phone', '')} | {personal.get('location', '')}"
    p = doc.add_paragraph(contact_info)
    p.alignment = 1 # Center
    
    # Summary
    if data.get('summary'):
        doc.add_heading('Professional Summary', level=1)
        doc.add_paragraph(data['summary'])

    # Experience
    if data.get('experience'):
        doc.add_heading('Work Experience', level=1)
        for exp in data['experience']:
            p = doc.add_paragraph()
            run = p.add_run(f"{exp.get('role', '')} at {exp.get('company', '')}")
            run.bold = True
            p.add_run(f"\n{exp.get('startDate', '')} - {exp.get('endDate', 'Present')}")
            doc.add_paragraph(exp.get('description', ''))

    # Education
    if data.get('education'):
        doc.add_heading('Education', level=1)
        for edu in data['education']:
            p = doc.add_paragraph()
            run = p.add_run(f"{edu.get('degree', '')}, {edu.get('school', '')}")
            run.bold = True
            p.add_run(f"\n{edu.get('startDate', '')} - {edu.get('endDate', '')}")

    # Skills
    if data.get('skills'):
        doc.add_heading('Skills', level=1)
        doc.add_paragraph(", ".join(data['skills']))

    doc.save(output_path)
    return True, "Word Document Generated Successfully"
